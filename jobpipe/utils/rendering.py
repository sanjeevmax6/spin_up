from __future__ import annotations

import logging
import re
from pathlib import Path

from docx import Document
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas

logger = logging.getLogger(__name__)


def slugify(value: str) -> str:
    value = value.strip().lower()
    value = re.sub(r"[^a-z0-9]+", "-", value)
    return value.strip("-") or "item"


def write_markdown(path: Path, content: str) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")
    logger.info("wrote markdown: %s", path)
    return path


def write_docx(path: Path, content: str) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    for line in content.splitlines():
        if line.strip():
            doc.add_paragraph(line)
        else:
            doc.add_paragraph("")
    doc.save(path)
    logger.info("wrote docx: %s", path)
    return path


def write_pdf(path: Path, content: str) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    c = canvas.Canvas(str(path), pagesize=LETTER)
    _, height = LETTER
    margin = inch
    y = height - margin
    c.setFont("Helvetica", 11)
    for line in content.splitlines():
        line = line.rstrip()
        if not line:
            y -= 14
        else:
            wrapped = _wrap_line(line, max_chars=95)
            for segment in wrapped:
                if y < margin:
                    c.showPage()
                    c.setFont("Helvetica", 11)
                    y = height - margin
                c.drawString(margin, y, segment)
                y -= 14
        if y < margin:
            c.showPage()
            c.setFont("Helvetica", 11)
            y = height - margin
    c.save()
    logger.info("wrote pdf: %s", path)
    return path


def _wrap_line(text: str, max_chars: int) -> list[str]:
    words = text.split()
    if not words:
        return [""]
    lines: list[str] = []
    current = words[0]
    for word in words[1:]:
        candidate = f"{current} {word}"
        if len(candidate) <= max_chars:
            current = candidate
        else:
            lines.append(current)
            current = word
    lines.append(current)
    return lines
