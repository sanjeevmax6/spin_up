from __future__ import annotations

import json
import logging
import re
import shlex
import subprocess
import time
from pathlib import Path
from typing import Any, Callable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from pydantic import ValidationError

from jobpipe.models import JobRow
from jobpipe.rules import evaluate_row_rules
from jobpipe.sheets import fetch_job_description
from jobpipe.state import GraphState
from jobpipe.utils.prompting import load_prompt, render_prompt
from jobpipe.utils.rendering import write_docx, write_markdown, write_pdf
from jobpipe.utils.rows import validate_required_columns

logger = logging.getLogger(__name__)


def ingest_row_node(state: GraphState) -> dict[str, Any]:
    node = "ingest_row_node"
    logger.info("[%s] start row=%s", node, state.get("row_number"))
    raw_row = dict(state.get("raw_row", {}))
    row_number = int(state.get("row_number", 0))
    missing = validate_required_columns(raw_row)
    if missing:
        return _mark_failed(state, node, f"missing required columns: {', '.join(missing)}")
    try:
        row = JobRow.from_sheet_row(row_number=row_number, raw=raw_row)
    except ValidationError as exc:
        msgs = [err["msg"] for err in exc.errors()]
        return _mark_failed(state, node, f"validation error: {'; '.join(msgs)}")
    updates = _mark_succeeded(state, node)
    updates["row_context"] = row
    return updates


def validate_context_node(state: GraphState) -> dict[str, Any]:
    node = "validate_context_node"
    logger.info("[%s] start row=%s", node, state.get("row_number"))
    row = state.get("row_context")
    app_config = state.get("app_config")
    if row is None or app_config is None:
        return _mark_failed(state, node, "missing row context or app config")

    if not row.job_description_text and row.job_description_url:
        try:
            row.job_description_text = fetch_job_description(row.job_description_url)
        except Exception as exc:  # noqa: BLE001
            return _mark_failed(state, node, f"unable to fetch job description URL: {exc}")

    rules = state.get("rules")
    if rules is None:
        return _mark_failed(state, node, "rules not loaded")

    rule_errors = evaluate_row_rules(row, rules)
    if rule_errors:
        return _mark_failed(state, node, "; ".join(rule_errors))

    resume_context = _load_yaml_file(app_config.resume_context_path)
    if not resume_context:
        return _mark_failed(
            state,
            node,
            f"resume context not found or empty: {app_config.resume_context_path}",
        )
    restrictions = _load_yaml_file(app_config.restrictions_path)

    updates = _mark_succeeded(state, node)
    updates["resume_context"] = resume_context
    updates["restrictions"] = restrictions
    updates["row_context"] = row
    return updates


def resume_node(state: GraphState) -> dict[str, Any]:
    node = "resume_node"
    logger.info("[%s] start row=%s", node, state.get("row_number"))
    app_config = state.get("app_config")
    row = state.get("row_context")
    llm_client = state.get("llm_client")
    resume_context = state.get("resume_context", {})
    restrictions = state.get("restrictions", {})
    if app_config is None or row is None or llm_client is None:
        return _mark_failed(state, node, "missing graph context for resume generation")

    try:
        shared_restrictions = load_prompt(app_config.prompt_dir, "shared_restrictions.prompt.md")
        selection_template = load_prompt(app_config.prompt_dir, "resume_selection.prompt.md")
        rewrite_template = load_prompt(app_config.prompt_dir, "resume_rewrite.prompt.md")
    except Exception as exc:  # noqa: BLE001
        return _mark_failed(state, node, f"unable to load resume templates: {exc}")

    dynamic_resume_context = {
        "experiences": resume_context.get("experiences", []),
        "projects": resume_context.get("projects", []),
        "constraints": resume_context.get("constraints", []),
    }

    selection_prompt = render_prompt(
        selection_template,
        {
            "shared_restrictions": shared_restrictions,
            "row_context": json.dumps(row.model_dump(), indent=2),
            "resume_context": json.dumps(dynamic_resume_context, indent=2),
            "extra_restrictions": json.dumps(restrictions, indent=2),
        },
    )
    try:
        selection_text, usage_1 = llm_client.complete(selection_prompt)
        logger.info("[%s] selection completed tokens=%s", node, usage_1.get("total_tokens", 0))
    except Exception as exc:  # noqa: BLE001
        return _mark_failed(state, node, f"resume selection failed: {exc}")

    selected_experience_ids, selected_project_ids = _extract_selected_ids(
        selection_text, dynamic_resume_context
    )
    selected_context = _select_context(
        dynamic_resume_context, selected_experience_ids, selected_project_ids
    )
    selected_ids_json = {
        "selected_experience_ids": selected_experience_ids,
        "selected_project_ids": selected_project_ids,
    }

    rewrite_prompt = render_prompt(
        rewrite_template,
        {
            "shared_restrictions": shared_restrictions,
            "row_context": json.dumps(row.model_dump(), indent=2),
            "selected_context": json.dumps(selected_context, indent=2),
            "extra_restrictions": json.dumps(restrictions, indent=2),
            "experience_bullet_max_words": int(app_config.experience_bullet_max_words),
            "project_bullet_max_words": int(app_config.project_bullet_max_words),
        },
    )

    try:
        rewritten_text, usage_2 = llm_client.complete(rewrite_prompt)
        logger.info("[%s] rewrite completed tokens=%s", node, usage_2.get("total_tokens", 0))
    except Exception as exc:  # noqa: BLE001
        return _mark_failed(state, node, f"resume rewrite failed: {exc}")

    try:
        refined_resume_json = _parse_refined_json(rewritten_text)
        refined_resume_json = _validate_refined_resume(
            refined_resume_json=refined_resume_json,
            selected_ids=selected_ids_json,
            selected_context=selected_context,
            exp_bullet_max_words=int(app_config.experience_bullet_max_words),
            proj_bullet_max_words=int(app_config.project_bullet_max_words),
            max_experiences=int(app_config.max_experiences),
            max_projects=int(app_config.max_projects_initial),
        )
    except Exception as exc:  # noqa: BLE001
        return _mark_failed(state, node, f"refined resume JSON validation failed: {exc}")

    merged_usage = {
        "input_tokens": int(usage_1.get("input_tokens", 0)) + int(usage_2.get("input_tokens", 0)),
        "output_tokens": int(usage_1.get("output_tokens", 0)) + int(usage_2.get("output_tokens", 0)),
        "total_tokens": int(usage_1.get("total_tokens", 0)) + int(usage_2.get("total_tokens", 0)),
    }

    return _mark_succeeded(
        state,
        node,
        artifacts={
            "selected_ids_json": selected_ids_json,
            "refined_resume_json": refined_resume_json,
        },
        usage=merged_usage,
    )


def cover_letter_node(state: GraphState) -> dict[str, Any]:
    node = "cover_letter_node"
    logger.info("[%s] start row=%s", node, state.get("row_number"))
    app_config = state.get("app_config")
    row = state.get("row_context")
    llm_client = state.get("llm_client")
    artifacts = dict(state.get("artifacts", {}))
    restrictions = state.get("restrictions", {})
    if app_config is None or row is None or llm_client is None:
        return _mark_failed(state, node, "missing graph context for cover letter")

    refined_resume_json = artifacts.get("refined_resume_json")
    if not isinstance(refined_resume_json, dict):
        return _mark_failed(state, node, "refined resume JSON missing; cannot build cover letter")

    try:
        shared_restrictions = load_prompt(app_config.prompt_dir, "shared_restrictions.prompt.md")
        template = load_prompt(app_config.prompt_dir, "cover_letter.prompt.md")
    except Exception as exc:  # noqa: BLE001
        return _mark_failed(state, node, f"unable to load cover letter template: {exc}")

    refined_summary = {
        "experiences": [
            {
                "id": item.get("id", ""),
                "company": item.get("company", ""),
                "role": item.get("role", ""),
                "bullets": item.get("bullets", []),
            }
            for item in refined_resume_json.get("experiences", [])
        ],
        "projects": [
            {
                "id": item.get("id", ""),
                "name": item.get("name", ""),
                "bullets": item.get("bullets", []),
            }
            for item in refined_resume_json.get("projects", [])
        ],
    }

    prompt = render_prompt(
        template,
        {
            "shared_restrictions": shared_restrictions,
            "row_context": json.dumps(row.model_dump(), indent=2),
            "refined_resume_json": json.dumps(refined_summary, indent=2),
            "extra_restrictions": json.dumps(restrictions, indent=2),
        },
    )
    try:
        text, usage = _run_with_retry(
            lambda: llm_client.complete(prompt),
            retries=int(app_config.non_resume_node_retries),
            backoff_seconds=float(app_config.non_resume_retry_backoff_seconds),
        )
    except Exception as exc:  # noqa: BLE001
        return _mark_failed(state, node, f"cover letter generation failed: {exc}")
    return _mark_succeeded(state, node, artifacts={"cover_letter_md": text}, usage=usage)


def linkedin_search_node(state: GraphState) -> dict[str, Any]:
    node = "linkedin_search_node"
    logger.info("[%s] start row=%s", node, state.get("row_number"))
    app_config = state.get("app_config")
    row = state.get("row_context")
    llm_client = state.get("llm_client")
    restrictions = state.get("restrictions", {})
    if app_config is None or row is None or llm_client is None:
        return _mark_failed(state, node, "missing graph context for linkedin search")

    try:
        shared_restrictions = load_prompt(app_config.prompt_dir, "shared_restrictions.prompt.md")
        template = load_prompt(app_config.prompt_dir, "linkedin_search.prompt.md")
    except Exception as exc:  # noqa: BLE001
        return _mark_failed(state, node, f"unable to load linkedin search template: {exc}")

    prompt = render_prompt(
        template,
        {
            "shared_restrictions": shared_restrictions,
            "row_context": json.dumps(row.model_dump(), indent=2),
            "extra_restrictions": json.dumps(restrictions, indent=2),
        },
    )
    try:
        text, usage = _run_with_retry(
            lambda: llm_client.complete(prompt),
            retries=int(app_config.non_resume_node_retries),
            backoff_seconds=float(app_config.non_resume_retry_backoff_seconds),
        )
    except Exception as exc:  # noqa: BLE001
        return _mark_failed(state, node, f"linkedin target generation failed: {exc}")
    return _mark_succeeded(state, node, artifacts={"linkedin_targets_md": text}, usage=usage)


def outreach_node(state: GraphState) -> dict[str, Any]:
    node = "outreach_node"
    logger.info("[%s] start row=%s", node, state.get("row_number"))
    app_config = state.get("app_config")
    row = state.get("row_context")
    llm_client = state.get("llm_client")
    restrictions = state.get("restrictions", {})
    artifacts = dict(state.get("artifacts", {}))
    statuses = dict(state.get("node_status", {}))
    if statuses.get("linkedin_search_node") != "succeeded":
        return _mark_skipped(state, node, "linkedin targets unavailable")
    if app_config is None or row is None or llm_client is None:
        return _mark_failed(state, node, "missing graph context for outreach")

    try:
        shared_restrictions = load_prompt(app_config.prompt_dir, "shared_restrictions.prompt.md")
        template = load_prompt(app_config.prompt_dir, "outreach.prompt.md")
    except Exception as exc:  # noqa: BLE001
        return _mark_failed(state, node, f"unable to load outreach template: {exc}")

    prompt = render_prompt(
        template,
        {
            "shared_restrictions": shared_restrictions,
            "row_context": json.dumps(row.model_dump(), indent=2),
            "linkedin_targets_md": artifacts.get("linkedin_targets_md", ""),
            "extra_restrictions": json.dumps(restrictions, indent=2),
        },
    )
    try:
        text, usage = _run_with_retry(
            lambda: llm_client.complete(prompt),
            retries=int(app_config.non_resume_node_retries),
            backoff_seconds=float(app_config.non_resume_retry_backoff_seconds),
        )
    except Exception as exc:  # noqa: BLE001
        return _mark_failed(state, node, f"outreach generation failed: {exc}")

    linkedin_notes, cold_email = _split_outreach_output(text)
    return _mark_succeeded(
        state,
        node,
        artifacts={"linkedin_notes_md": linkedin_notes, "cold_email_md": cold_email},
        usage=usage,
    )


def render_node(state: GraphState) -> dict[str, Any]:
    node = "render_node"
    logger.info("[%s] start row=%s", node, state.get("row_number"))
    app_config = state.get("app_config")
    output_dir = state.get("output_dir")
    artifacts = dict(state.get("artifacts", {}))
    if not output_dir or app_config is None:
        return _mark_failed(state, node, "output_dir or app config not set")

    base = Path(output_dir)
    base.mkdir(parents=True, exist_ok=True)
    file_artifacts: dict[str, str] = {}
    trim_log: list[str] = []
    trimmed_resume_obj: dict[str, Any] | None = None

    try:
        selected_ids_json = artifacts.get("selected_ids_json")
        if isinstance(selected_ids_json, dict):
            selected_path = base / "selected_ids.json"
            selected_path.write_text(json.dumps(selected_ids_json, indent=2), encoding="utf-8")
            file_artifacts["selected_ids_file"] = selected_path.name

        refined_resume_json = artifacts.get("refined_resume_json")
        if isinstance(refined_resume_json, dict):
            logger.info("[%s] rendering resume artifacts", node)
            refined_path = base / "refined_resume.json"
            refined_path.write_text(json.dumps(refined_resume_json, indent=2), encoding="utf-8")
            file_artifacts["refined_resume_file"] = refined_path.name

            trimmed_resume_json, trim_log = _apply_one_page_policy(refined_resume_json, app_config)
            trimmed_path = base / "trimmed_resume.json"
            trimmed_path.write_text(json.dumps(trimmed_resume_json, indent=2), encoding="utf-8")
            file_artifacts["trimmed_resume_file"] = trimmed_path.name
            trimmed_resume_obj = trimmed_resume_json

            resume_docx_path = base / "resume.docx"
            _render_resume_docx(trimmed_resume_json, app_config, resume_docx_path)
            file_artifacts["resume_docx"] = resume_docx_path.name
            logger.info("[%s] wrote %s", node, resume_docx_path)

            resume_pdf_path = base / "resume.pdf"
            pdf_ok, pdf_error = _convert_docx_to_pdf(resume_docx_path, resume_pdf_path, app_config)
            if pdf_ok:
                file_artifacts["resume_pdf"] = resume_pdf_path.name
                logger.info("[%s] wrote %s", node, resume_pdf_path)
            else:
                trim_log.append(f"pdf_conversion_warning: {pdf_error}")
                logger.warning("[%s] pdf conversion warning: %s", node, pdf_error)

        if artifacts.get("cover_letter_md"):
            cover_path = write_markdown(base / "cover_letter.md", artifacts["cover_letter_md"])
            file_artifacts["cover_letter_md"] = str(cover_path.name)
            file_artifacts["cover_letter_docx"] = str(
                write_docx(base / "cover_letter.docx", artifacts["cover_letter_md"]).name
            )
            file_artifacts["cover_letter_pdf"] = str(
                write_pdf(base / "cover_letter.pdf", artifacts["cover_letter_md"]).name
            )
        if artifacts.get("linkedin_targets_md"):
            file_artifacts["linkedin_targets_md"] = str(
                write_markdown(base / "linkedin_targets.md", artifacts["linkedin_targets_md"]).name
            )
        if artifacts.get("linkedin_notes_md"):
            file_artifacts["linkedin_notes_md"] = str(
                write_markdown(base / "linkedin_connection_notes.md", artifacts["linkedin_notes_md"]).name
            )
        if artifacts.get("cold_email_md"):
            file_artifacts["cold_email_md"] = str(
                write_markdown(base / "cold_email.md", artifacts["cold_email_md"]).name
            )
    except Exception as exc:  # noqa: BLE001
        logger.exception("[%s] render failed", node)
        return _mark_failed(state, node, f"render failed: {exc}")

    merged_artifacts = dict(file_artifacts)
    if isinstance(artifacts.get("selected_ids_json"), dict):
        merged_artifacts["selected_ids_json"] = artifacts["selected_ids_json"]
    if isinstance(artifacts.get("refined_resume_json"), dict):
        merged_artifacts["refined_resume_json"] = artifacts["refined_resume_json"]
    if trimmed_resume_obj is not None:
        merged_artifacts["trimmed_resume_json"] = trimmed_resume_obj
    elif isinstance(artifacts.get("refined_resume_json"), dict):
        merged_artifacts["trimmed_resume_json"] = artifacts["refined_resume_json"]
    merged_artifacts["trim_log"] = trim_log
    return _mark_succeeded(state, node, artifacts=merged_artifacts)


def report_node(state: GraphState) -> dict[str, Any]:
    node = "report_node"
    logger.info("[%s] start row=%s", node, state.get("row_number"))
    output_dir = state.get("output_dir")
    if not output_dir:
        return _mark_succeeded(state, node)
    base = Path(output_dir)
    base.mkdir(parents=True, exist_ok=True)

    artifacts = dict(state.get("artifacts", {}))
    trimmed = artifacts.get("trimmed_resume_json") if isinstance(artifacts.get("trimmed_resume_json"), dict) else {}

    manifest = {
        "row_number": state.get("row_number"),
        "node_status": state.get("node_status", {}),
        "node_errors": state.get("node_errors", {}),
        "token_usage_by_node": state.get("token_usage_by_node", {}),
        "artifacts": artifacts,
        "resume_layout": {
            "experience_count": len(trimmed.get("experiences", [])) if isinstance(trimmed, dict) else 0,
            "project_count": len(trimmed.get("projects", [])) if isinstance(trimmed, dict) else 0,
            "trim_log": artifacts.get("trim_log", []),
        },
    }
    manifest_path = base / "manifest.json"
    manifest_path.write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    updates = _mark_succeeded(state, node)
    merged_artifacts = dict(updates.get("artifacts", {}))
    merged_artifacts["manifest_json"] = manifest_path.name
    updates["artifacts"] = merged_artifacts
    return updates


# ---------- Shared helpers ----------
def _ensure_maps(
    state: GraphState,
) -> tuple[dict[str, str], dict[str, list[str]], dict[str, dict[str, int]], dict[str, Any]]:
    node_status = dict(state.get("node_status", {}))
    node_errors = dict(state.get("node_errors", {}))
    token_usage_by_node = dict(state.get("token_usage_by_node", {}))
    artifacts = dict(state.get("artifacts", {}))
    return node_status, node_errors, token_usage_by_node, artifacts


def _mark_succeeded(
    state: GraphState,
    node: str,
    artifacts: dict[str, Any] | None = None,
    usage: dict[str, int] | None = None,
) -> dict[str, Any]:
    node_status, node_errors, token_usage_by_node, merged_artifacts = _ensure_maps(state)
    node_status[node] = "succeeded"
    node_errors.setdefault(node, [])
    if artifacts:
        merged_artifacts.update(artifacts)
    if usage:
        token_usage_by_node[node] = {
            "input_tokens": int(usage.get("input_tokens", 0)),
            "output_tokens": int(usage.get("output_tokens", 0)),
            "total_tokens": int(usage.get("total_tokens", 0)),
        }
    logger.info("[%s] status=succeeded", node)
    return {
        "node_status": node_status,
        "node_errors": node_errors,
        "token_usage_by_node": token_usage_by_node,
        "artifacts": merged_artifacts,
    }


def _mark_failed(
    state: GraphState, node: str, error: str, usage: dict[str, int] | None = None
) -> dict[str, Any]:
    node_status, node_errors, token_usage_by_node, merged_artifacts = _ensure_maps(state)
    node_status[node] = "failed"
    node_errors.setdefault(node, []).append(error)
    logger.error("[%s] status=failed error=%s", node, error)
    if usage:
        token_usage_by_node[node] = {
            "input_tokens": int(usage.get("input_tokens", 0)),
            "output_tokens": int(usage.get("output_tokens", 0)),
            "total_tokens": int(usage.get("total_tokens", 0)),
        }
    return {
        "node_status": node_status,
        "node_errors": node_errors,
        "token_usage_by_node": token_usage_by_node,
        "artifacts": merged_artifacts,
    }


def _mark_skipped(state: GraphState, node: str, reason: str) -> dict[str, Any]:
    node_status, node_errors, token_usage_by_node, merged_artifacts = _ensure_maps(state)
    node_status[node] = "skipped"
    node_errors.setdefault(node, []).append(reason)
    logger.warning("[%s] status=skipped reason=%s", node, reason)
    return {
        "node_status": node_status,
        "node_errors": node_errors,
        "token_usage_by_node": token_usage_by_node,
        "artifacts": merged_artifacts,
    }


def _load_yaml_file(path: str | Path) -> dict[str, Any]:
    import yaml

    file_path = Path(path)
    if not file_path.exists():
        return {}
    with file_path.open("r", encoding="utf-8") as f:
        data = yaml.safe_load(f) or {}
    if not isinstance(data, dict):
        return {}
    return data


def _run_with_retry(
    fn: Callable[[], tuple[str, dict[str, int]]],
    retries: int,
    backoff_seconds: float,
) -> tuple[str, dict[str, int]]:
    attempts = max(0, retries) + 1
    last_error: Exception | None = None
    for attempt in range(attempts):
        try:
            if attempt > 0:
                logger.info("retrying operation attempt=%s/%s", attempt + 1, attempts)
            return fn()
        except Exception as exc:  # noqa: BLE001
            last_error = exc
            if attempt == attempts - 1:
                break
            time.sleep(max(0.0, backoff_seconds * (2**attempt)))
    assert last_error is not None
    raise last_error


# ---------- Resume JSON helpers ----------
def _extract_selected_ids(
    selection_text: str, resume_context: dict[str, Any]
) -> tuple[list[str], list[str]]:
    exp_ids = [str(item.get("id", "")).strip() for item in resume_context.get("experiences", []) if item.get("id")]
    proj_ids = [str(item.get("id", "")).strip() for item in resume_context.get("projects", []) if item.get("id")]

    parsed = _parse_json_from_text(selection_text)
    selected_exp: list[str] = []
    selected_proj: list[str] = []
    if isinstance(parsed, dict):
        selected_exp = [str(x).strip() for x in parsed.get("selected_experience_ids", []) if str(x).strip()]
        selected_proj = [str(x).strip() for x in parsed.get("selected_project_ids", []) if str(x).strip()]

    if not selected_exp or not selected_proj:
        lower = selection_text.lower()
        if not selected_exp:
            selected_exp = [item for item in exp_ids if item.lower() in lower]
        if not selected_proj:
            selected_proj = [item for item in proj_ids if item.lower() in lower]

    selected_exp = [x for x in selected_exp if x in exp_ids]
    selected_proj = [x for x in selected_proj if x in proj_ids]

    if not selected_exp:
        selected_exp = exp_ids[: min(3, len(exp_ids))]
    if not selected_proj:
        selected_proj = proj_ids[: min(4, len(proj_ids))]

    return list(dict.fromkeys(selected_exp))[:3], list(dict.fromkeys(selected_proj))[:4]


def _select_context(
    resume_context: dict[str, Any],
    selected_experience_ids: list[str],
    selected_project_ids: list[str],
) -> dict[str, Any]:
    experiences = resume_context.get("experiences", [])
    projects = resume_context.get("projects", [])

    selected_experiences = [e for e in experiences if str(e.get("id", "")) in selected_experience_ids]
    selected_projects = [p for p in projects if str(p.get("id", "")) in selected_project_ids]
    return {
        "experiences": selected_experiences,
        "projects": selected_projects,
    }


def _parse_refined_json(text: str) -> dict[str, Any]:
    parsed = _parse_json_from_text(text)
    if not isinstance(parsed, dict):
        raise ValueError("LLM output is not valid JSON object")
    return parsed


def _parse_json_from_text(text: str) -> Any:
    stripped = text.strip()
    try:
        return json.loads(stripped)
    except Exception:  # noqa: BLE001
        pass

    fence_match = re.search(r"```(?:json)?\s*(\{[\s\S]*\})\s*```", text, flags=re.IGNORECASE)
    if fence_match:
        return json.loads(fence_match.group(1))

    # Fallback: first JSON object span.
    start = text.find("{")
    end = text.rfind("}")
    if start != -1 and end != -1 and end > start:
        return json.loads(text[start : end + 1])
    raise ValueError("No JSON object found in text")


def _validate_refined_resume(
    refined_resume_json: dict[str, Any],
    selected_ids: dict[str, list[str]],
    selected_context: dict[str, Any],
    exp_bullet_max_words: int,
    proj_bullet_max_words: int,
    max_experiences: int,
    max_projects: int,
) -> dict[str, Any]:
    selected_exp_ids = set(selected_ids.get("selected_experience_ids", []))
    selected_proj_ids = set(selected_ids.get("selected_project_ids", []))

    source_exp_map = {
        str(item.get("id", "")): item for item in selected_context.get("experiences", [])
    }
    source_proj_map = {
        str(item.get("id", "")): item for item in selected_context.get("projects", [])
    }

    experiences_out: list[dict[str, Any]] = []
    projects_out: list[dict[str, Any]] = []

    for exp in refined_resume_json.get("experiences", []):
        exp_id = str(exp.get("id", "")).strip()
        if exp_id not in selected_exp_ids or exp_id not in source_exp_map:
            raise ValueError(f"unknown or unselected experience id: {exp_id}")
        src = source_exp_map[exp_id]
        bullets = [str(x).strip() for x in exp.get("bullets", []) if str(x).strip()]
        if len(bullets) < 3:
            src_bullets = [str(x).strip() for x in src.get("bullets", []) if str(x).strip()]
            while len(bullets) < 3 and src_bullets:
                bullets.append(src_bullets[len(bullets) % len(src_bullets)])
        bullets = bullets[:3]
        bullets = [_trim_words(b, exp_bullet_max_words) for b in bullets]

        tech_text = str(exp.get("tech_bullet", "")).strip()
        if not tech_text:
            tech_src = src.get("tech", [])
            if isinstance(tech_src, list):
                tech_text = "Tech: " + ", ".join([str(x).strip() for x in tech_src if str(x).strip()][:8])
            else:
                tech_text = "Tech: Relevant stack from source resume"
        if not tech_text.lower().startswith("tech:"):
            tech_text = "Tech: " + tech_text
        tech_text = _trim_words(tech_text, exp_bullet_max_words)
        bullets.append(tech_text)

        experiences_out.append(
            {
                "id": exp_id,
                "company": str(src.get("company", "")).strip(),
                "role": str(src.get("role", "")).strip(),
                "location": str(src.get("location", "")).strip(),
                "date_range": str(src.get("date_range", "")).strip(),
                "bullets": bullets,
            }
        )

    for proj in refined_resume_json.get("projects", []):
        proj_id = str(proj.get("id", "")).strip()
        if proj_id not in selected_proj_ids or proj_id not in source_proj_map:
            raise ValueError(f"unknown or unselected project id: {proj_id}")
        src = source_proj_map[proj_id]

        bullets = [str(x).strip() for x in proj.get("bullets", []) if str(x).strip()]
        bullets = _drop_code_bullets(bullets)
        src_bullets = [str(x).strip() for x in src.get("bullets", []) if str(x).strip()]
        while len(bullets) < 2 and src_bullets:
            bullets.append(src_bullets[len(bullets) % len(src_bullets)])
        bullets = bullets[:3]
        bullets = [_trim_words(b, proj_bullet_max_words) for b in bullets]

        code_link = str(proj.get("code_link", "")).strip()
        if not code_link:
            links = src.get("links", [])
            if isinstance(links, list) and links:
                code_link = str(links[0]).strip()

        projects_out.append(
            {
                "id": proj_id,
                "name": str(src.get("name", "")).strip(),
                "bullets": bullets,
                "code_link": code_link,
                "techs": _normalize_project_techs(src.get("tech", [])),
            }
        )

    if not experiences_out:
        for exp_id in list(selected_exp_ids)[:max_experiences]:
            src = source_exp_map[exp_id]
            src_bullets = [str(x).strip() for x in src.get("bullets", []) if str(x).strip()][:3]
            tech_src = src.get("tech", [])
            tech_bullet = "Tech: " + ", ".join([str(x).strip() for x in tech_src][:8]) if isinstance(tech_src, list) else "Tech:"
            experiences_out.append(
                {
                    "id": exp_id,
                    "company": str(src.get("company", "")).strip(),
                    "role": str(src.get("role", "")).strip(),
                    "location": str(src.get("location", "")).strip(),
                    "date_range": str(src.get("date_range", "")).strip(),
                    "bullets": [_trim_words(x, exp_bullet_max_words) for x in src_bullets] + [_trim_words(tech_bullet, exp_bullet_max_words)],
                }
            )

    if not projects_out:
        for proj_id in list(selected_proj_ids)[:max_projects]:
            src = source_proj_map[proj_id]
            links = src.get("links", [])
            projects_out.append(
                {
                    "id": proj_id,
                    "name": str(src.get("name", "")).strip(),
                    "bullets": [_trim_words(str(x).strip(), proj_bullet_max_words) for x in src.get("bullets", [])[:3]],
                    "code_link": str(links[0]).strip() if isinstance(links, list) and links else "",
                    "techs": _normalize_project_techs(src.get("tech", [])),
                }
            )

    experiences_out = experiences_out[:max_experiences]
    projects_out = projects_out[:max_projects]

    for exp in experiences_out:
        exp["bullets"] = (exp.get("bullets", []) + ["Tech:"])[:4]
        if len(exp["bullets"]) < 4:
            exp["bullets"] = exp["bullets"] + ["Tech:"] * (4 - len(exp["bullets"]))
        if not str(exp["bullets"][3]).lower().startswith("tech:"):
            exp["bullets"][3] = "Tech: " + str(exp["bullets"][3])
        src = source_exp_map.get(str(exp.get("id", "")).strip(), {})
        src_bullets = [str(x).strip() for x in src.get("bullets", []) if str(x).strip()]
        for i in range(3):
            if _word_count(exp["bullets"][i]) < 6 and i < len(src_bullets):
                exp["bullets"][i] = _trim_words(src_bullets[i], exp_bullet_max_words)

    for proj in projects_out:
        proj_bullets = [str(x).strip() for x in proj.get("bullets", []) if str(x).strip()]
        proj_bullets = _drop_code_bullets(proj_bullets)
        if len(proj_bullets) < 2:
            proj_bullets = (proj_bullets + ["Relevant project implementation detail"] * 2)[:2]
        proj["bullets"] = proj_bullets[:3]
        src = source_proj_map.get(str(proj.get("id", "")).strip(), {})
        src_bullets = [str(x).strip() for x in src.get("bullets", []) if str(x).strip()]
        for i in range(len(proj["bullets"])):
            if _word_count(proj["bullets"][i]) < 6 and i < len(src_bullets):
                proj["bullets"][i] = _trim_words(src_bullets[i], proj_bullet_max_words)

    return {"experiences": experiences_out, "projects": projects_out}


def _trim_words(text: str, max_words: int) -> str:
    text = str(text).strip()
    words = text.split()
    if len(words) <= max_words:
        return text
    trimmed = " ".join(words[:max_words]).strip().rstrip(",;:-")
    if trimmed and trimmed[-1] not in ".!?":
        trimmed += "."
    return trimmed


def _apply_one_page_policy(refined: dict[str, Any], app_config: Any) -> tuple[dict[str, Any], list[str]]:
    experiences = list(refined.get("experiences", []))[: int(app_config.max_experiences)]
    projects = list(refined.get("projects", []))[: int(app_config.max_projects_initial)]
    log: list[str] = []

    exp_word_cap = int(app_config.experience_bullet_max_words)
    proj_word_cap = int(app_config.project_bullet_max_words)

    for exp in experiences:
        exp["bullets"] = [_trim_words(str(b), exp_word_cap) for b in exp.get("bullets", [])][:4]
        if len(exp["bullets"]) < 4:
            exp["bullets"] += ["Tech:"] * (4 - len(exp["bullets"]))
        if not str(exp["bullets"][3]).lower().startswith("tech:"):
            exp["bullets"][3] = "Tech: " + str(exp["bullets"][3])

    for proj in projects:
        proj["bullets"] = [_trim_words(str(b), proj_word_cap) for b in proj.get("bullets", [])][:3]
        if len(proj["bullets"]) < 2:
            proj["bullets"] += ["Relevant project detail"] * (2 - len(proj["bullets"]))

    target_score = _layout_score_limit()
    while _layout_score(experiences, projects) > target_score and len(projects) > int(app_config.min_projects):
        removed = projects.pop()
        log.append(f"trim_projects_count: removed {removed.get('id', '')}")

    while _layout_score(experiences, projects) > target_score:
        changed = False
        for proj in projects:
            bullets = proj.get("bullets", [])
            if len(bullets) > 2:
                proj["bullets"] = bullets[:2]
                log.append(f"trim_project_bullets_to_2: {proj.get('id', '')}")
                changed = True
                if _layout_score(experiences, projects) <= target_score:
                    break
        if not changed:
            break

    return {"experiences": experiences, "projects": projects}, log


def _layout_score(experiences: list[dict[str, Any]], projects: list[dict[str, Any]]) -> int:
    score = 0
    score += 2  # heading area
    for exp in experiences:
        score += 2  # company/role line + spacing
        score += len(exp.get("bullets", []))
    for proj in projects:
        score += 1
        score += len(proj.get("bullets", []))
        if str(proj.get("code_link", "")).strip():
            score += 1
    return score


def _layout_score_limit() -> int:
    # Heuristic target for one-page docx template with compact spacing.
    return 34


def _ensure_resume_template(template_path: Path) -> None:
    if template_path.exists():
        return
    template_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading("EXPERIENCE", level=1)
    doc.add_paragraph("{% for exp in experiences %}")
    doc.add_paragraph("{{ exp.company }} | {{ exp.role }} | {{ exp.location }} | {{ exp.date_range }}")
    doc.add_paragraph("- {{ exp.bullets[0] }}")
    doc.add_paragraph("- {{ exp.bullets[1] }}")
    doc.add_paragraph("- {{ exp.bullets[2] }}")
    doc.add_paragraph("- {{ exp.bullets[3] }}")
    doc.add_paragraph("{% endfor %}")

    doc.add_heading("PROJECTS", level=1)
    doc.add_paragraph("{% for proj in projects %}")
    doc.add_paragraph("{{ proj.name }}")
    doc.add_paragraph("- {{ proj.bullets[0] }}")
    doc.add_paragraph("- {{ proj.bullets[1] }}")
    doc.add_paragraph("{% if proj.bullets|length > 2 %}- {{ proj.bullets[2] }}{% endif %}")
    doc.add_paragraph("Code: {{ proj.code_link }}")
    doc.add_paragraph("{% endfor %}")
    doc.save(template_path)


def _render_resume_docx(trimmed_resume_json: dict[str, Any], app_config: Any, output_docx_path: Path) -> None:
    template_path = Path(str(app_config.resume_template_path))
    _ensure_resume_template(template_path)
    payload = _build_docxtpl_payload(trimmed_resume_json)
    doc = Document(str(template_path))

    work_heading = _find_heading_paragraph(doc, ["WORK EXPERIENCE", "WORKEXPERIENCE", "EXPERIENCE"])
    proj_heading = _find_heading_paragraph(doc, ["PROJECTS", "PROJECT"])
    tech_heading = _find_heading_paragraph(doc, ["TECHNICAL SKILLS", "SKILLS"])
    if tech_heading is None:
        tech_heading = _find_heading_paragraph(doc, ["EDUCATION"])
    if work_heading is None or proj_heading is None:
        sample = [
            p.text.strip()
            for p in doc.paragraphs
            if p.text and p.text.strip()
        ][:40]
        logger.warning(
            "headings not found in template; falling back to docxtpl rendering work=%s projects=%s paragraphs_sample=%s",
            work_heading is not None,
            proj_heading is not None,
            sample,
        )
        _render_resume_docx_docxtpl(template_path, payload, output_docx_path)
        return

    logger.info(
        "rendering structured resume template experiences=%s projects=%s",
        len(trimmed_resume_json.get("experiences", [])),
        len(trimmed_resume_json.get("projects", [])),
    )
    _clear_between(doc, work_heading, proj_heading)
    if tech_heading is not None:
        _clear_between(doc, proj_heading, tech_heading)
    else:
        _clear_after(doc, proj_heading)

    for exp in trimmed_resume_json.get("experiences", []):
        _insert_experience_before(doc, proj_heading, exp)

    anchor = tech_heading
    for proj in trimmed_resume_json.get("projects", []):
        if anchor is not None:
            _insert_project_before(doc, anchor, proj)
        else:
            _append_project(doc, proj)

    _cleanup_doc_whitespace(doc)
    output_docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx_path))


def _build_docxtpl_payload(trimmed_resume_json: dict[str, Any]) -> dict[str, Any]:
    experiences_payload: list[dict[str, Any]] = []
    for exp in trimmed_resume_json.get("experiences", []):
        bullets = [str(b).strip() for b in exp.get("bullets", []) if str(b).strip()]
        while len(bullets) < 4:
            bullets.append("Tech:")
        bullets = bullets[:4]
        bullets[3] = _as_technologies_line(bullets[3])
        experiences_payload.append(
            {
                "id": str(exp.get("id", "")).strip(),
                "company": str(exp.get("company", "")).strip(),
                "role": str(exp.get("role", "")).strip(),
                "location": str(exp.get("location", "")).strip(),
                "date_range": str(exp.get("date_range", "")).strip(),
                "bullets": bullets,
            }
        )

    projects_payload: list[dict[str, Any]] = []
    for proj in trimmed_resume_json.get("projects", []):
        bullets = [str(b).strip() for b in proj.get("bullets", []) if str(b).strip()][:3]
        while len(bullets) < 2:
            bullets.append("Relevant project detail")
        techs = [str(t).strip() for t in proj.get("techs", []) if str(t).strip()][:6]
        projects_payload.append(
            {
                "id": str(proj.get("id", "")).strip(),
                "name": str(proj.get("name", "")).strip(),
                "bullets": bullets,
                "code_link": str(proj.get("code_link", "")).strip(),
                "techs": techs,
                "techs_bracket": f"[{', '.join(techs)}]" if techs else "",
            }
        )
    return {"experiences": experiences_payload, "projects": projects_payload}


def _render_resume_docx_docxtpl(
    template_path: Path, payload: dict[str, Any], output_docx_path: Path
) -> None:
    try:
        from docxtpl import DocxTemplate
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError(
            "docxtpl is required for placeholder rendering. Install dependencies with `pip install -e .`."
        ) from exc
    logger.info(
        "rendering docxtpl fallback template experiences=%s projects=%s template=%s",
        len(payload.get("experiences", [])),
        len(payload.get("projects", [])),
        template_path,
    )
    doc = DocxTemplate(str(template_path))
    doc.render(payload)
    output_docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx_path))


def _convert_docx_to_pdf(input_docx: Path, output_pdf: Path, app_config: Any) -> tuple[bool, str]:
    cmd_template = str(getattr(app_config, "pdf_converter_command", "") or "").strip()
    if cmd_template:
        cmd = cmd_template.format(input=shlex.quote(str(input_docx)), output=shlex.quote(str(output_pdf)))
        try:
            subprocess.run(cmd, shell=True, check=True, capture_output=True, text=True)
            return output_pdf.exists(), ""
        except Exception as exc:  # noqa: BLE001
            return False, str(exc)

    soffice = subprocess.run(["which", "soffice"], capture_output=True, text=True)
    if soffice.returncode == 0 and soffice.stdout.strip():
        outdir = str(output_pdf.parent)
        try:
            subprocess.run(
                [
                    soffice.stdout.strip(),
                    "--headless",
                    "--convert-to",
                    "pdf",
                    "--outdir",
                    outdir,
                    str(input_docx),
                ],
                check=True,
                capture_output=True,
                text=True,
            )
            converted = output_pdf.parent / (input_docx.stem + ".pdf")
            if converted.exists() and converted != output_pdf:
                converted.replace(output_pdf)
            return output_pdf.exists(), ""
        except Exception as exc:  # noqa: BLE001
            return False, str(exc)

    return False, "No PDF converter configured (set pdf_converter_command)"


def _split_outreach_output(text: str) -> tuple[str, str]:
    marker_1 = "## LinkedIn Connection Notes"
    marker_2 = "## Cold Email"
    if marker_1 in text and marker_2 in text:
        left, right = text.split(marker_2, 1)
        linkedin = left.strip()
        cold = (marker_2 + "\n" + right.strip()).strip()
        return linkedin + "\n", cold + "\n"
    return text.strip() + "\n", text.strip() + "\n"


def _normalize_project_techs(raw: Any) -> list[str]:
    if isinstance(raw, list):
        vals = [str(x).strip() for x in raw if str(x).strip()]
    else:
        text = str(raw or "").strip()
        vals = [x.strip() for x in re.split(r"[,\|/]+", text) if x.strip()]
    return vals[:6]


def _find_paragraph_by_text(doc: Document, text: str) -> Any:
    target = text.strip().lower()
    for p in doc.paragraphs:
        normalized = " ".join(p.text.replace("\u00a0", " ").strip().lower().split())
        if normalized == target or target in normalized:
            return p
    return None


def _find_heading_paragraph(doc: Document, candidates: list[str]) -> Any:
    normalized_candidates = {_normalize_heading_text(c) for c in candidates}
    for p in doc.paragraphs:
        normalized = _normalize_heading_text(p.text)
        if normalized in normalized_candidates:
            return p
    return None


def _normalize_heading_text(text: str) -> str:
    return "".join(ch for ch in text.upper() if ch.isalnum())


def _clear_between(doc: Document, start_para: Any, end_para: Any) -> None:
    started = False
    to_delete: list[Any] = []
    for p in doc.paragraphs:
        if p == start_para:
            started = True
            continue
        if p == end_para:
            break
        if started:
            to_delete.append(p)
    for p in to_delete:
        p._element.getparent().remove(p._element)


def _clear_after(doc: Document, start_para: Any) -> None:
    started = False
    to_delete: list[Any] = []
    for p in doc.paragraphs:
        if p == start_para:
            started = True
            continue
        if started:
            to_delete.append(p)
    for p in to_delete:
        p._element.getparent().remove(p._element)


def _insert_experience_before(doc: Document, anchor_para: Any, exp: dict[str, Any]) -> None:
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = False
    tbl.columns[0].width = Inches(5.1)
    tbl.columns[1].width = Inches(2.0)
    tbl.style = "Table Grid"
    _remove_table_borders(tbl)

    left_p = tbl.cell(0, 0).paragraphs[0]
    role = _compact_role(str(exp.get("role", "")).strip(), max_words=7)
    r = left_p.add_run(str(exp.get("company", "")).strip())
    r.bold = True
    left_p.add_run(" | ")
    left_p.add_run(role)
    _tighten_paragraph(left_p)

    right_p = tbl.cell(0, 1).paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right_p.add_run(_compact_location(str(exp.get("location", "")).strip(), max_words=5))
    right_p.add_run(" | ")
    rd = right_p.add_run(str(exp.get("date_range", "")).strip())
    rd.bold = True
    _tighten_paragraph(right_p)

    anchor_para._p.addprevious(tbl._tbl)

    bullets = [str(b).strip() for b in exp.get("bullets", []) if str(b).strip()]
    for i, b in enumerate(bullets[:4]):
        p = _insert_bullet_paragraph_before(anchor_para)
        run = p.add_run(b if i < 3 else _as_technologies_line(b))
        if i == 3:
            run.italic = True
            run.bold = False
        _tighten_paragraph(p)


def _insert_project_before(doc: Document, anchor_para: Any, proj: dict[str, Any]) -> None:
    tbl = doc.add_table(rows=1, cols=2)
    tbl.autofit = False
    tbl.columns[0].width = Inches(5.3)
    tbl.columns[1].width = Inches(1.8)
    tbl.style = "Table Grid"
    _remove_table_borders(tbl)

    left_p = tbl.cell(0, 0).paragraphs[0]
    rn = left_p.add_run(str(proj.get("name", "")).strip())
    rn.bold = True

    techs = [str(t).strip() for t in proj.get("techs", []) if str(t).strip()][:6]
    if techs:
        rt = left_p.add_run(f" [{', '.join(techs)}]")
        rt.italic = True
        rt.bold = False

    right_p = tbl.cell(0, 1).paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    link = str(proj.get("code_link", "")).strip()
    if link:
        _add_hyperlink(right_p, link, "Code")
    else:
        right_p.add_run("Code")

    anchor_para._p.addprevious(tbl._tbl)

    bullets = [str(b).strip() for b in proj.get("bullets", []) if str(b).strip()][:3]
    for b in bullets:
        p = _insert_bullet_paragraph_before(anchor_para)
        p.add_run(b)
        _tighten_paragraph(p)


def _append_resume_content(doc: Document, trimmed_resume_json: dict[str, Any]) -> None:
    doc.add_paragraph("WORK EXPERIENCE")
    for exp in trimmed_resume_json.get("experiences", []):
        p = doc.add_paragraph()
        p.add_run(str(exp.get("company", ""))).bold = True
        p.add_run(" | ")
        p.add_run(str(exp.get("role", "")))
        p.add_run(" | ")
        p.add_run(str(exp.get("location", "")))
        p.add_run(" | ")
        p.add_run(str(exp.get("date_range", ""))).bold = True
        for i, b in enumerate(exp.get("bullets", [])[:4]):
            bp = _append_bullet_paragraph(doc)
            rb = bp.add_run(_as_technologies_line(str(b)) if i == 3 else str(b))
            if i == 3:
                rb.italic = True
    doc.add_paragraph("PROJECTS")
    for proj in trimmed_resume_json.get("projects", []):
        p = doc.add_paragraph()
        p.add_run(str(proj.get("name", ""))).bold = True
        techs = [str(t).strip() for t in proj.get("techs", []) if str(t).strip()][:6]
        if techs:
            r = p.add_run(f" [{', '.join(techs)}]")
            r.italic = True
        p.add_run("  Code")
        for b in proj.get("bullets", [])[:3]:
            bp = _append_bullet_paragraph(doc)
            bp.add_run(str(b))


def _as_technologies_line(text: str) -> str:
    t = str(text).strip()
    if t.lower().startswith("tech:"):
        return "Technologies:" + t[5:]
    if t.lower().startswith("technologies:"):
        return t
    return "Technologies: " + t


def _remove_table_borders(table: Any) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "nil")
        borders.append(elem)
    tblPr.append(borders)


def _tighten_paragraph(p: Any) -> None:
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.0


def _append_project(doc: Document, proj: dict[str, Any]) -> None:
    p = doc.add_paragraph()
    p.add_run(str(proj.get("name", ""))).bold = True
    techs = [str(t).strip() for t in proj.get("techs", []) if str(t).strip()][:6]
    if techs:
        r = p.add_run(f" [{', '.join(techs)}]")
        r.italic = True
    for b in [str(x).strip() for x in proj.get("bullets", []) if str(x).strip()][:3]:
        bp = _append_bullet_paragraph(doc)
        bp.add_run(b)


def _add_hyperlink(paragraph: Any, url: str, text: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)
    text_elem = OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _insert_bullet_paragraph_before(anchor_para: Any) -> Any:
    p = anchor_para.insert_paragraph_before("")
    if not _try_set_list_bullet_style(p):
        p.add_run("• ")
    return p


def _append_bullet_paragraph(doc: Document) -> Any:
    p = doc.add_paragraph("")
    if not _try_set_list_bullet_style(p):
        p.add_run("• ")
    return p


def _try_set_list_bullet_style(paragraph: Any) -> bool:
    try:
        paragraph.style = "List Bullet"
        return True
    except Exception:  # noqa: BLE001
        logger.debug("List Bullet style missing in template; using literal bullet prefix")
        return False


def _cleanup_doc_whitespace(doc: Document) -> None:
    to_delete: list[Any] = []
    prev_blank = False
    for p in doc.paragraphs:
        txt = p.text.replace("\u00a0", " ").strip()
        is_template = ("{{" in txt) or ("{%" in txt) or ("{#" in txt) or ("#}" in txt)
        is_bullet_only = txt in {"•", "o", "-", "*"}
        if is_template or is_bullet_only:
            to_delete.append(p)
            continue
        is_blank = txt == ""
        if is_blank and prev_blank:
            to_delete.append(p)
            continue
        prev_blank = is_blank
    for p in to_delete:
        p._element.getparent().remove(p._element)


def _compact_role(role: str, max_words: int = 7) -> str:
    tokens = [w for w in role.split() if w]
    if len(tokens) <= max_words:
        return role
    return " ".join(tokens[:max_words]).rstrip(",;:-")


def _compact_location(location: str, max_words: int = 5) -> str:
    tokens = [w for w in location.split() if w]
    if len(tokens) <= max_words:
        return location
    return " ".join(tokens[:max_words]).rstrip(",;:-")


def _word_count(text: str) -> int:
    return len([w for w in str(text).split() if w.strip()])


def _drop_code_bullets(bullets: list[str]) -> list[str]:
    kept: list[str] = []
    for b in bullets:
        low = b.lower()
        has_url = ("http://" in low) or ("https://" in low) or ("github.com/" in low)
        looks_like_code_line = low.startswith("[code") or low.startswith("code:")
        if has_url or looks_like_code_line:
            continue
        kept.append(b)
    return kept
